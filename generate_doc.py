from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# ── Styles ──
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)
    hs.font.name = 'Calibri'

# ── Title Page ──
doc.add_paragraph('')
doc.add_paragraph('')
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Vendr — Ecommerce Application')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Technical Design & Implementation Document')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph('')
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run('Version 2.0  •  April 2026\n').font.size = Pt(11)
meta.add_run('Technology Stack: React · TypeScript · Vite · Tailwind CSS v4').font.size = Pt(11)

doc.add_page_break()

# ── Table of Contents ──
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Project Overview',
    '2. Technology Stack',
    '3. Project Structure',
    '4. Design System & Theme',
    '5. Core Features',
    '6. Page-by-Page Breakdown',
    '7. State Management',
    '8. New Features Added',
    '9. Copywriting & Brand Voice',
    '10. Build & Deployment',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(4)

doc.add_page_break()

# ═══════════════════════════════════════════════
# 1. Project Overview
# ═══════════════════════════════════════════════
doc.add_heading('1. Project Overview', level=1)
doc.add_paragraph(
    'Vendr is a global ecommerce web application that connects shoppers with authentic, '
    'handpicked products from artisans and sellers in 14+ countries. The app features a '
    'modern dark-themed UI with red accent colors, responsive layouts, and a rich set of '
    'interactive features including a shopping cart, wishlist, country-based filtering, '
    'product comparison, deals, and more.'
)
doc.add_paragraph(
    'The application was originally scaffolded from a Figma-to-code export, then underwent '
    'two major design passes: (1) a full layout polish to improve the UI/UX shell, and '
    '(2) a complete black-and-red rebrand across every component and page. A third pass '
    'added new features, fresh copywriting, and entirely new page layouts.'
)

# ═══════════════════════════════════════════════
# 2. Technology Stack
# ═══════════════════════════════════════════════
doc.add_heading('2. Technology Stack', level=1)

table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = 'Layer'
hdr[1].text = 'Technology'
hdr[2].text = 'Purpose'
for cell in hdr:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

data = [
    ('Framework', 'React 19 + TypeScript', 'Component-based UI with type safety'),
    ('Build Tool', 'Vite 6.3', 'Lightning-fast dev server & HMR'),
    ('Styling', 'Tailwind CSS v4', 'Utility-first CSS with @theme inline tokens'),
    ('Routing', 'React Router v7', 'Client-side SPA navigation'),
    ('Icons', 'Lucide React', 'Consistent, tree-shakeable icon library'),
    ('UI Primitives', 'Radix UI + shadcn/ui', 'Accessible, headless UI components'),
    ('Fonts', 'Manrope + Fraunces', 'Modern sans-serif + elegant display serif'),
    ('Mobile', 'Capacitor', 'iOS & Android native wrappers'),
    ('Animations', 'tw-animate-css + custom keyframes', 'Fade-in, slide, marquee, pulse'),
]
for layer, tech, purpose in data:
    row = table.add_row().cells
    row[0].text = layer
    row[1].text = tech
    row[2].text = purpose

# ═══════════════════════════════════════════════
# 3. Project Structure
# ═══════════════════════════════════════════════
doc.add_heading('3. Project Structure', level=1)
doc.add_paragraph(
    'The codebase follows a feature-oriented folder structure inside src/app/:'
)

structure = [
    ('src/app/pages/', 'Route-level page components (Home, Products, Cart, About, Contact, etc.)'),
    ('src/app/components/', 'Shared components (Header, Footer, ProductCard, Chatbot, etc.)'),
    ('src/app/components/ui/', 'shadcn/ui primitives (Button, Dialog, Card, etc.)'),
    ('src/app/context/', 'React Context providers (CartContext, CountryContext, WishlistContext)'),
    ('src/app/data/', 'Static product data (products.ts)'),
    ('src/styles/', 'Global CSS: theme.css (design tokens), tailwind.css (animations), fonts.css'),
]
for folder, desc in structure:
    p = doc.add_paragraph()
    run = p.add_run(folder)
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = 'Consolas'
    p.add_run(f'  —  {desc}')

# ═══════════════════════════════════════════════
# 4. Design System & Theme
# ═══════════════════════════════════════════════
doc.add_heading('4. Design System & Theme', level=1)

doc.add_heading('4.1 Color Palette', level=2)
doc.add_paragraph(
    'The entire app uses a dark-mode-first palette defined via CSS custom properties '
    'in theme.css, then mapped to Tailwind tokens via a @theme inline block.'
)

color_table = doc.add_table(rows=1, cols=3)
color_table.style = 'Light Grid Accent 1'
ch = color_table.rows[0].cells
ch[0].text = 'Token'
ch[1].text = 'Value'
ch[2].text = 'Usage'
for cell in ch:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

colors = [
    ('--background', '#0a0a0a', 'Page backgrounds'),
    ('--foreground', '#fafafa', 'Primary text color'),
    ('--primary / --accent', '#dc2626 (Red 600)', 'Buttons, links, badges, focus rings'),
    ('--card', '#141414', 'Card surfaces, modals, panels'),
    ('--border', 'rgba(255,255,255,0.08)', 'Subtle dividers'),
    ('--muted', '#1a1a1a / #a1a1a1', 'Secondary surfaces / muted text'),
]
for token, value, usage in colors:
    row = color_table.add_row().cells
    row[0].text = token
    row[1].text = value
    row[2].text = usage

doc.add_heading('4.2 Typography', level=2)
doc.add_paragraph(
    'Manrope is used for all body text and UI elements. Fraunces is available as a display/serif '
    'font for decorative headings. Weights used: 400 (regular), 600 (semibold), 700 (bold), '
    '900 (black). Headings are typically "font-black" (900) with tight tracking.'
)

doc.add_heading('4.3 Spacing & Layout Conventions', level=2)
doc.add_paragraph(
    '• Max content width: max-w-7xl (80rem / 1280px)\n'
    '• Section padding: py-24 (6rem vertical)\n'
    '• Card radius: rounded-2xl (1rem) or rounded-3xl (1.5rem)\n'
    '• Card border: border border-white/5 with hover:border-red-600/20\n'
    '• Standard gap: gap-6 for grids, gap-4 for tighter layouts'
)

doc.add_heading('4.4 Animation', level=2)
doc.add_paragraph(
    'Custom keyframes are defined in tailwind.css:\n'
    '• fade-in: opacity 0→1 with translateY for staggered card reveals\n'
    '• marquee: infinite horizontal scroll for the country ticker\n'
    '• float: gentle vertical bob for decorative elements\n'
    '• Native Tailwind: animate-pulse, animate-ping, animate-bounce for micro-interactions'
)

# ═══════════════════════════════════════════════
# 5. Core Features
# ═══════════════════════════════════════════════
doc.add_heading('5. Core Features', level=1)

features = [
    ('Shopping Cart', 'Add/remove/update quantities. Persisted via React Context. Cart count displayed as a badge in the header.'),
    ('Wishlist', 'Heart-toggle on every product card feeds into WishlistContext. Dedicated /wishlist page shows saved items with "Move to Cart" action.'),
    ('Country Filtering', 'Global country selector in the header (CountryContext) automatically filters products across pages. Local per-page country filter also available.'),
    ('Product Catalog', '17+ products across 6 categories (Electronics, Fashion, Home, Beauty, Food, Accessories) from 14 countries, each with images, ratings, reviews, and descriptions.'),
    ('Sort & View Toggle', 'Products page offers sorting by price (asc/desc), rating, and name, plus a grid/list view toggle.'),
    ('Deals & Promotions', 'Dedicated /deals page with timed promotional cards and newsletter subscription.'),
    ('Responsive Design', 'Fully responsive from mobile (1 col) through tablet (2 col) to desktop (4 col grids). Mobile hamburger menu with slide-in navigation.'),
    ('Scroll-Aware Header', 'Header hides on scroll-down and reappears on scroll-up for maximum content visibility.'),
    ('Chatbot', 'Floating chat widget for customer support (Chatbot.tsx).'),
]
for title, desc in features:
    p = doc.add_paragraph()
    run = p.add_run(f'{title}:  ')
    run.bold = True
    p.add_run(desc)

# ═══════════════════════════════════════════════
# 6. Page-by-Page Breakdown
# ═══════════════════════════════════════════════
doc.add_heading('6. Page-by-Page Breakdown', level=1)

doc.add_heading('6.1 Home Page  (/)', level=2)
doc.add_paragraph(
    'The homepage is the most content-rich page, designed to immediately communicate brand '
    'identity and drive product discovery. It consists of 7 distinct sections:'
)
home_sections = [
    ('Hero (Split Layout)', 'Left column: bold headline ("Not your average store."), subtext, Shop Now / Today\'s Deals CTAs, social-proof avatars with star rating. Right column: 2×2 product image collage with country labels and hover zoom. Floating badge shows "14 Countries • 17+ Products".'),
    ('Country Marquee', 'Infinite horizontally-scrolling ticker of all 14 country flags + names. Uses CSS @keyframes marquee animation with duplicated content for seamless loop.'),
    ('Category Bento Grid', 'Asymmetric grid of 5 category cards (Electronics spans 2×2, Beauty spans 1×2). Each card has a background image, icon, item count, and hover-reveal arrow. Links to /products.'),
    ('Trending Products', '4-card grid of featured products using the shared ProductCard component. Staggered fade-in animation.'),
    ('New Arrivals (Horizontal Scroll)', 'Horizontally-scrollable row of 6 product cards with snap scrolling. Allows quick browsing without leaving the page.'),
    ('Value Props Strip', '4-column strip with icons: Free Global Shipping, Authenticity Guaranteed, Delivered in Days, Hassle-Free Returns. Divided by subtle vertical lines.'),
    ('Testimonials', '2×2 grid of customer review cards with star ratings, quote text, avatar initials, and location. Quote icon as decorative element.'),
    ('Newsletter CTA', 'Centered section with email input + subscribe button. Soft red glow background. "Get the good stuff first." headline.'),
]
for title, desc in home_sections:
    p = doc.add_paragraph()
    run = p.add_run(f'{title}:  ')
    run.bold = True
    p.add_run(desc)

doc.add_heading('6.2 Products Page  (/products)', level=2)
doc.add_paragraph(
    'Full catalog browser with filtering, sorting, and dual view modes.\n\n'
    '• Header: "The good stuff. All in one place." with product/country counts.\n'
    '• Filter panel: Category pill buttons + Country pill buttons (with flags). '
    'Active filter shown in red-600. Global country override notice when active.\n'
    '• Toolbar: Product count badge, sort dropdown (Default / Price Low→High / Price High→Low / '
    'Top Rated / A→Z), grid/list view toggle buttons.\n'
    '• Grid view: Responsive 1–4 column card grid.\n'
    '• List view: Horizontal cards with image, metadata, description, price, and rating inline.\n'
    '• Empty state: Illustration, message, Reset Filters + Back to Home buttons.'
)

doc.add_heading('6.3 Contact Page  (/contact)', level=2)
doc.add_paragraph(
    'Two-panel layout designed for conversion:\n\n'
    '• Left column (2/5 width): Stacked info cards — Email, Phone, Live Chat, Visit Us — '
    'each with icon that fills red on hover. Plus a Business Hours card with day/time table.\n'
    '• Right column (3/5 width): Contact form with Name + Email side-by-side, Subject dropdown '
    '(6 predefined topics), Message textarea, and full-width Send button.\n'
    '• FAQ Section: 4 collapsible accordion questions with chevron rotation animation. '
    'Covers shipping, returns, authenticity, and international delivery.'
)

doc.add_heading('6.4 About Page  (/about)', level=2)
doc.add_paragraph(
    'Brand narrative page with visual storytelling:\n\n'
    '• Hero: Left-aligned asymmetric layout. "We built the store we wanted to shop at." '
    'Red gradient on second line.\n'
    '• Stats Strip: 4-column bar — 2M+ customers, 14 countries, 150+ destinations, 4.9 rating.\n'
    '• Mission Section: Two-column — left has heading, right has 3 value cards (Authenticity, '
    'Artisans, Speed) with icon hover effects.\n'
    '• Timeline: Vertical timeline from 2019–2024 with red dots on a line, showing company milestones.\n'
    '• CTA: "Ready to see what the fuss is about?" with Shop button.'
)

doc.add_heading('6.5 Deals Page  (/deals)', level=2)
doc.add_paragraph(
    'Promotional landing page:\n\n'
    '• Hero: "Deals that don\'t wait around." with Flame badge.\n'
    '• Deal Cards: 2×2 grid. Each card shows discount amount (large type), title, description, '
    'expiry, and Shop button. Flash Sale marked with animated HOT badge.\n'
    '• Newsletter: "Never miss a deal." email signup.'
)

doc.add_heading('6.6 Wishlist Page  (/wishlist)', level=2)
doc.add_paragraph(
    'Dedicated wishlist management:\n\n'
    '• Shows saved item count. Each item displayed as a card with image, country, name, price.\n'
    '• Two actions per item: "Add to Cart" (moves item to cart and removes from wishlist) '
    'and delete (removes from wishlist).\n'
    '• Empty state: Heart icon, message, "Explore Products" CTA.'
)

doc.add_heading('6.7 Other Pages', level=2)
doc.add_paragraph(
    '• Cart (/cart): Full cart management with quantity controls, item removal, order summary.\n'
    '• Checkout (/checkout): Multi-step checkout form.\n'
    '• Product Detail (/products/:id): Full product page with image, description, add-to-cart.\n'
    '• Pricing (/pricing): Pricing chart comparison.\n'
    '• Comparison (/comparison): Side-by-side product comparison.\n'
    '• Help (/help): FAQ and support resources.\n'
    '• 404 (/*): Custom not-found page.'
)

# ═══════════════════════════════════════════════
# 7. State Management
# ═══════════════════════════════════════════════
doc.add_heading('7. State Management', level=1)
doc.add_paragraph(
    'The app uses React Context for global state, avoiding external state libraries:'
)
ctx_table = doc.add_table(rows=1, cols=3)
ctx_table.style = 'Light Grid Accent 1'
ch = ctx_table.rows[0].cells
ch[0].text = 'Context'
ch[1].text = 'File'
ch[2].text = 'Provides'
for cell in ch:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

ctxs = [
    ('CartContext', 'context/CartContext.tsx', 'cart[], addToCart, removeFromCart, updateQuantity, clearCart, getCartTotal, getCartCount'),
    ('WishlistContext', 'context/WishlistContext.tsx', 'wishlist[], addToWishlist, removeFromWishlist, isInWishlist, getWishlistCount'),
    ('CountryContext', 'context/CountryContext.tsx', 'selectedCountry, setSelectedCountry, getCurrency — enables global country filtering and currency display'),
]
for name, file, provides in ctxs:
    row = ctx_table.add_row().cells
    row[0].text = name
    row[1].text = file
    row[2].text = provides

doc.add_paragraph(
    '\nAll three providers wrap the RouterProvider in App.tsx, ensuring every page '
    'and component has access to cart, wishlist, and country state.'
)

# ═══════════════════════════════════════════════
# 8. New Features Added
# ═══════════════════════════════════════════════
doc.add_heading('8. New Features Added', level=1)
doc.add_paragraph(
    'The following features and improvements were added during the latest design iteration:'
)

new_features = [
    ('Functional Wishlist System', 'Created WishlistContext with add/remove/check operations. Wired the heart button on every ProductCard to this context (replacing the previous local state toggle). Added a dedicated /wishlist route and page. Header heart icon now shows a live count badge.'),
    ('Newsletter Signup', 'Email input + Subscribe button embedded in the Home page footer section. Also present on the Deals page. Uses controlled form state with submission feedback.'),
    ('Customer Testimonials', 'Four review cards on the Home page with 5-star ratings, quoted text, customer initials as avatars, and geographic locations. Builds social proof.'),
    ('FAQ Accordion', 'Interactive collapsible question/answer section on the Contact page. Uses local state to track which question is open. Chevron icon rotates with CSS transition.'),
    ('Product Sorting', 'Dropdown on the Products page allowing sort by: Default, Price (Low→High), Price (High→Low), Top Rated, Alphabetical (A→Z). Implemented via Array.sort() on the filtered product list.'),
    ('Grid/List View Toggle', 'Two-button toggle on the Products page switches between the standard card grid and a new horizontal list layout showing image + full details inline.'),
    ('Category Bento Grid', 'Asymmetric image grid on the Home page showcasing 5 product categories with item counts, icons, and hover effects. Uses CSS Grid with col-span/row-span for layout variety.'),
    ('Country Marquee Ticker', 'Infinite horizontal scroll of all 14 country names + flags at the top of the Home page. Implemented with CSS @keyframes marquee and duplicated content nodes.'),
    ('Horizontal-Scroll New Arrivals', 'Snap-scrolling horizontal product row on the Home page showing 6 recently added items. Uses overflow-x-auto with snap-x snap-mandatory.'),
    ('Company Timeline', 'Vertical timeline on the About page showing 5 milestones from 2019 to 2024, with red-dot markers and a connecting vertical line.'),
]
for title, desc in new_features:
    p = doc.add_paragraph()
    run = p.add_run(f'{title}:  ')
    run.bold = True
    p.add_run(desc)

# ═══════════════════════════════════════════════
# 9. Copywriting & Brand Voice
# ═══════════════════════════════════════════════
doc.add_heading('9. Copywriting & Brand Voice', level=1)
doc.add_paragraph(
    'All page copy was rewritten to shift from generic marketing language to a direct, '
    'confident, and conversational tone. The goal: sound like a friend who runs a cool store, '
    'not a corporate brochure.'
)

voice_table = doc.add_table(rows=1, cols=2)
voice_table.style = 'Light Grid Accent 1'
vh = voice_table.rows[0].cells
vh[0].text = 'Before'
vh[1].text = 'After'
for cell in vh:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

voices = [
    ('"Discover Global Treasures"', '"Not your average store."'),
    ('"Experience Luxury Shopping Redefined"', '"Find your thing."'),
    ('"Discover Unique Treasures Worldwide"', '"The good stuff. All in one place."'),
    ('"Contact Us — We\'d love to hear from you"', '"We\'re all ears."'),
    ('"About Vendr"', '"We built the store we wanted to shop at."'),
    ('"Unbeatable Deals"', '"Deals that don\'t wait around."'),
    ('"Featured Collection: Trending Worldwide"', '"Selling fast. Don\'t blink."'),
    ('"Never Miss a Deal"', '"Never miss a deal." (kept — already good)'),
]
for before, after in voices:
    row = voice_table.add_row().cells
    row[0].text = before
    row[1].text = after

# ═══════════════════════════════════════════════
# 10. Build & Deployment
# ═══════════════════════════════════════════════
doc.add_heading('10. Build & Deployment', level=1)

doc.add_heading('10.1 Development', level=2)
doc.add_paragraph(
    'npm run dev — Starts Vite dev server with hot module replacement (HMR).\n'
    'The app runs on http://localhost:5173/ (or the next available port).\n'
    'All file changes are reflected instantly in the browser.'
)

doc.add_heading('10.2 Production Build', level=2)
doc.add_paragraph(
    'npm run build — Produces optimized static assets in the dist/ folder.\n'
    'Vite handles tree-shaking, code splitting, CSS minification, and asset hashing.\n'
    'Output can be deployed to any static hosting (Vercel, Netlify, Cloudflare Pages, S3, etc.).'
)

doc.add_heading('10.3 Mobile (Capacitor)', level=2)
doc.add_paragraph(
    'The project includes capacitor.config.json for wrapping the web app as a native '
    'iOS and Android application. After building, use npx cap sync to copy web assets '
    'into native projects.'
)

# ── Save ──
output_path = r'c:\Users\rahul\Downloads\Ecommerce app\Vendr_Project_Documentation.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
